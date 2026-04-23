"""File extraction and upload handling services."""

import mimetypes
import os
import re
import shutil
import stat
import xml.etree.ElementTree as ET
import zipfile

import requests
from shared.ai.llm_client import call_ai
from urllib.parse import unquote, urlparse

from apps.qq_ai_bridge.adapters.napcat_client import (
    fetch_napcat_file_download_info,
    get_msg_detail,
    send_group_msg,
    send_private_msg,
)
from apps.qq_ai_bridge.config.settings import (
    ALLOWED_PRIVATE_USER,
    BASE_DATA_DIR,
    GROUP_UPLOAD_DIR,
    MAX_ARCHIVE_LISTING,
    MAX_FILE_CONTENT_LEN,
    PRIVATE_UPLOAD_DIR,
    TEXT_LIKE_EXTS,
)
from storage_utils import load_private_context


def extract_file_info(event_data):
    """Extract file info from NapCat / OneBot message payloads."""
    def _first_present(payload: dict, *keys: str):
        for key in keys:
            value = payload.get(key)
            if value not in (None, "", []):
                return value
        return None

    def _normalize_file_info(payload: dict) -> dict:
        return {
            "name": _first_present(payload, "fileName", "name", "file_name", "file"),
            "url": _first_present(payload, "downloadUrl", "url", "fileUrl", "file_url"),
            "size": _first_present(payload, "fileSize", "size", "file_size"),
            "uuid": _first_present(payload, "fileUuid", "fileId", "file_id", "uuid"),
            "sub_id": _first_present(payload, "fileSubId", "subId", "file_sub_id", "sub_id"),
            "path": _first_present(payload, "filePath", "path", "file_path"),
            "raw": payload,
        }

    def _is_meaningful(info: dict) -> bool:
        return any(info.get(key) for key in ("name", "url", "uuid", "path"))

    def _parse_cq_file(raw_text: str) -> dict | None:
        # Example: [CQ:file,file=xxx.docx,url=https://...,file_id=...]
        if "[CQ:file" not in raw_text:
            return None
        match = re.search(r"\[CQ:file,([^\]]+)\]", raw_text)
        if not match:
            return None
        params = {}
        for item in match.group(1).split(","):
            if "=" not in item:
                continue
            key, value = item.split("=", 1)
            params[key.strip()] = value.strip()
        info = _normalize_file_info(params)
        return info if _is_meaningful(info) else None

    raw_message = event_data.get("message")
    raw_obj = event_data.get("raw", {})
    elements = raw_obj.get("elements", [])
    if not elements:
        elements = event_data.get("elements", [])
    file_signal_detected = False
    if isinstance(elements, list):
        for elem in elements:
            if not isinstance(elem, dict):
                continue

            file_elem = elem.get("fileElement")
            if file_elem is not None:
                file_signal_detected = True
            if isinstance(file_elem, dict):
                file_info = _normalize_file_info(file_elem)
                if _is_meaningful(file_info):
                    print(f"[FILE] extract_file_info 命中 raw.elements: {file_info}")
                    return file_info

    if isinstance(raw_message, list):
        for seg in raw_message:
            if not isinstance(seg, dict):
                continue
            if seg.get("type") == "file":
                file_signal_detected = True
                data = seg.get("data", {})
                if isinstance(data, dict):
                    file_info = _normalize_file_info(data)
                    if _is_meaningful(file_info):
                        print(f"[FILE] extract_file_info 命中 message.file: {file_info}")
                        return file_info
    elif isinstance(raw_message, str):
        file_signal_detected = file_signal_detected or ("[CQ:file" in raw_message)
        cq_info = _parse_cq_file(raw_message)
        if cq_info:
            print(f"[FILE] extract_file_info 命中 raw_message CQ:file: {cq_info}")
            return cq_info

    top_level_file = event_data.get("file")
    if isinstance(top_level_file, dict):
        file_signal_detected = True
        file_info = _normalize_file_info(top_level_file)
        if _is_meaningful(file_info):
            print(f"[FILE] extract_file_info 命中 top-level file: {file_info}")
            return file_info

    if file_signal_detected:
        print("[FILE] extract_file_info 检测到文件消息但未解析出完整文件信息")
    return None


def resolve_file_download_info(file_info):
    """Resolve file download URL from event or NapCat file API."""
    direct_url = file_info.get("url")
    if direct_url:
        print(f"[FILE_API] 使用事件自带 URL: {direct_url}")
        return direct_url, "url_from_event"

    resolved_url, reason = fetch_napcat_file_download_info(file_info)
    if resolved_url:
        file_info["url"] = resolved_url
        print(f"[FILE_API] 解析下载地址成功: {resolved_url}")
        return resolved_url, reason

    print(f"[FILE_API] 解析下载地址失败: {reason}")
    return None, reason


def safe_filename(name: str) -> str:
    """Convert a potentially unsafe filename into a safe local filename."""
    if not name:
        return "unknown_file"
    return name.replace("/", "_").replace("\\", "_").strip()


def derive_filename(file_info: dict) -> str | None:
    """Derive filename from known metadata when name is missing."""
    name = str(file_info.get("name") or "").strip()
    if name:
        return name

    file_url = str(file_info.get("url") or "").strip()
    if file_url:
        parsed = urlparse(file_url)
        tail = os.path.basename(parsed.path or "")
        tail = unquote(tail).strip()
        if tail:
            return tail

    file_path = str(file_info.get("path") or "").strip()
    if file_path:
        tail = os.path.basename(file_path).strip()
        if tail:
            return tail

    file_uuid = str(file_info.get("uuid") or "").strip()
    if file_uuid:
        return file_uuid
    return None


def describe_fs_entry(path):
    """Return a human-readable description of a filesystem entry."""
    try:
        st = os.stat(path)
        mode = stat.filemode(st.st_mode)
        return (
            f"path={path!r}, mode={mode}, uid={st.st_uid}, gid={st.st_gid}, "
            f"size={st.st_size}, readable={os.access(path, os.R_OK)}"
        )
    except Exception as e:
        return f"path={path!r}, stat_failed={e}"


def download_file_if_possible(file_info, save_dir):
    """Download or copy a file attachment into the target directory."""
    name = safe_filename(derive_filename(file_info))
    local_path = file_info.get("path")
    target_path = os.path.join(save_dir, name)
    url, resolve_reason = resolve_file_download_info(file_info)

    if url:
        if url.startswith("/app/.config/QQ"):
            try:
                host_path = url.replace("/app/.config/QQ", os.path.expanduser("~/napcat/qq"))
                print(f"[FILE] 检测到 NapCat 容器路径: {url} -> {host_path}")
                if not os.path.exists(host_path):
                    reason = f"napcat_host_path_missing: {host_path}"
                    print(f"[FILE] 容器文件不存在: {reason}")
                    return None, reason
                if not os.access(host_path, os.R_OK):
                    reason = f"napcat_host_path_not_readable: {describe_fs_entry(host_path)}"
                    print(f"[FILE] 容器文件无读取权限: {reason}")
                    return None, reason
                shutil.copy(host_path, target_path)
                print(f"[FILE] 容器文件复制成功: {target_path}")
                return target_path, "copied_from_napcat"
            except Exception as e:
                print(f"[FILE] 容器文件复制失败: {e}")
                return None, f"copy_from_napcat_failed: {e}"
        else:
            try:
                print(f"[FILE] 通过 URL 下载: {url} -> {target_path}")
                r = requests.get(url, timeout=30)
                r.raise_for_status()
                with open(target_path, "wb") as f:
                    f.write(r.content)
                print(f"[FILE] 下载成功: {target_path}")
                return target_path, "downloaded_by_url"
            except Exception as e:
                print(f"[FILE] URL 下载失败: {e}")
                return None, f"url_download_failed: {e}"

    print(f"[FILE] 未拿到可用下载链接，准备尝试本地路径。原因: {resolve_reason}")

    if local_path and os.path.exists(local_path):
        try:
            shutil.copy(local_path, target_path)
            print(f"[FILE] 本地复制成功: {target_path}")
            return target_path, "copied_from_local_path"
        except Exception as e:
            print(f"[FILE] 本地复制失败: {e}")
            return None, f"copy_local_path_failed: {e}"

    reason = f"no_download_url_and_local_path_unavailable: {local_path}"
    print(f"[FILE] 文件保存失败: {reason}, local_path={local_path!r}")
    return None, reason


def read_text_file(path):
    """Read a text file with fallback encodings."""
    for encoding in ("utf-8", "gb18030", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()[:MAX_FILE_CONTENT_LEN]
        except UnicodeDecodeError as e:
            print(f"[FILE] 读取失败(编码): {e}")
            continue
        except Exception as e:
            print(f"[FILE] 读取失败: {e}")
            return ""
    return ""


def extract_pdf_text(path):
    """Best-effort PDF text extraction with multiple backend fallbacks."""
    tried = []

    try:
        import fitz  # type: ignore

        doc = fitz.open(path)
        text = "\n".join(page.get_text() for page in doc).strip()
        doc.close()
        if text:
            print(f"[FILE] PDF 文本提取成功 (fitz): {path}")
            return text[:MAX_FILE_CONTENT_LEN], "pdf_text_fitz"
        tried.append("fitz:empty")
    except ImportError:
        tried.append("fitz:missing")
    except Exception as e:
        tried.append(f"fitz:err={e}")

    try:
        import pypdf  # type: ignore

        reader = pypdf.PdfReader(path)
        text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        if text:
            print(f"[FILE] PDF 文本提取成功 (pypdf): {path}")
            return text[:MAX_FILE_CONTENT_LEN], "pdf_text_pypdf"
        tried.append("pypdf:empty")
    except ImportError:
        tried.append("pypdf:missing")
    except Exception as e:
        tried.append(f"pypdf:err={e}")

    try:
        from pdfminer.high_level import extract_text as _pdfminer_extract  # type: ignore

        text = (_pdfminer_extract(path) or "").strip()
        if text:
            print(f"[FILE] PDF 文本提取成功 (pdfminer): {path}")
            return text[:MAX_FILE_CONTENT_LEN], "pdf_text_pdfminer"
        tried.append("pdfminer:empty")
    except ImportError:
        tried.append("pdfminer:missing")
    except Exception as e:
        tried.append(f"pdfminer:err={e}")

    print(f"[FILE] PDF 文本提取失败: tried={tried}")
    return None, None


_DOCX_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_docx_text(path):
    """Extract text from DOCX files, preferring python-docx when available."""
    try:
        import docx  # type: ignore

        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if text:
            print(f"[FILE] DOCX 文本提取成功 (python-docx): {path}")
            return text[:MAX_FILE_CONTENT_LEN], "docx_text_pydocx"
    except ImportError:
        pass
    except Exception as e:
        print(f"[FILE] DOCX (python-docx) 提取失败，回退到 zipfile 方案: {e}")

    try:
        with zipfile.ZipFile(path, "r") as zf:
            with zf.open("word/document.xml") as f:
                tree = ET.parse(f)
        texts = [node.text for node in tree.iter(f"{_DOCX_W_NS}t") if node.text]
        if not texts:
            texts = [node.text for node in tree.iter() if node.text]
        text = "\n".join(texts).strip()
        if text:
            print(f"[FILE] DOCX 文本提取成功 (zip/xml): {path}")
            return text[:MAX_FILE_CONTENT_LEN], "docx_text_xml"
    except Exception as e:
        print(f"[FILE] DOCX 文本提取失败: {e}")
    return None, None


def extract_pptx_text(path):
    """Extract text from PPTX files."""
    try:
        texts = []
        with zipfile.ZipFile(path, "r") as zf:
            for name in zf.namelist():
                if name.startswith("ppt/slides/slide") and name.endswith(".xml"):
                    with zf.open(name) as f:
                        tree = ET.parse(f)
                    texts.extend(node.text for node in tree.iter() if node.text)
        text = "\n".join(texts).strip()
        if text:
            print(f"[FILE] PPTX 文本提取成功: {path}")
            return text[:MAX_FILE_CONTENT_LEN], "pptx_text"
    except Exception as e:
        print(f"[FILE] PPTX 文本提取失败: {e}")
    return None, None


def extract_xlsx_text(path):
    """Extract text from XLSX shared strings and worksheet XML."""
    try:
        texts = []
        with zipfile.ZipFile(path, "r") as zf:
            for name in zf.namelist():
                if name.endswith(".xml") and (name.startswith("xl/sharedStrings") or name.startswith("xl/worksheets/")):
                    with zf.open(name) as f:
                        tree = ET.parse(f)
                    texts.extend(node.text for node in tree.iter() if node.text)
        text = "\n".join(texts).strip()
        if text:
            print(f"[FILE] XLSX 文本提取成功: {path}")
            return text[:MAX_FILE_CONTENT_LEN], "xlsx_text"
    except Exception as e:
        print(f"[FILE] XLSX 文本提取失败: {e}")
    return None, None


def extract_zip_summary(path):
    """Summarize a ZIP archive."""
    try:
        with zipfile.ZipFile(path, "r") as zf:
            names = zf.namelist()
        preview = names[:MAX_ARCHIVE_LISTING]
        summary = "ZIP 文件结构：\n" + "\n".join(preview)
        if len(names) > MAX_ARCHIVE_LISTING:
            summary += f"\n... 其余 {len(names) - MAX_ARCHIVE_LISTING} 个文件省略"
        print(f"[FILE] ZIP 结构提取成功: {path}")
        return summary[:MAX_FILE_CONTENT_LEN], "zip_summary"
    except Exception as e:
        print(f"[FILE] ZIP 结构提取失败: {e}")
    return None, None


def build_binary_file_summary(path, filename):
    """Build a metadata-only summary for binary files."""
    mime_type, _ = mimetypes.guess_type(filename or path)
    size = os.path.getsize(path)
    print(f"[FILE] 生成二进制文件摘要: {filename!r}, mime={mime_type!r}, size={size}")
    return (
        f"这是一个二进制文件。\n文件名：{filename}\nMIME：{mime_type}\n大小：{size} 字节"
    )[:MAX_FILE_CONTENT_LEN], "binary_summary"


DOC_LIKE_EXTS = {".pdf", ".docx", ".pptx", ".xlsx"}


def extract_file_content_for_ai(path, filename):
    """Extract best-effort readable content from an uploaded file."""
    ext = os.path.splitext(filename or path)[1].lower()
    mime_type, _ = mimetypes.guess_type(filename or path)
    print(f"[FILE] 开始提取文件内容: path={path}, ext={ext!r}, mime={mime_type!r}")

    if ext in TEXT_LIKE_EXTS or (mime_type and mime_type.startswith("text/")):
        content = read_text_file(path)
        if content:
            return content, "text_direct"

    if ext == ".pdf":
        content, reason = extract_pdf_text(path)
        if content:
            return content, reason
    if ext == ".docx":
        content, reason = extract_docx_text(path)
        if content:
            return content, reason
    if ext == ".pptx":
        content, reason = extract_pptx_text(path)
        if content:
            return content, reason
    if ext == ".xlsx":
        content, reason = extract_xlsx_text(path)
        if content:
            return content, reason

    if ext in DOC_LIKE_EXTS:
        hint = (
            f"文件 {filename} 是 {ext[1:].upper()} 文档，但本地解析失败，"
            "可能缺少对应的解析依赖（PyMuPDF / python-docx / openpyxl），"
            "或文档本身为扫描件/图片版。"
        )
        print(f"[FILE] 文档解析失败，返回缺失依赖提示: {hint}")
        return hint[:MAX_FILE_CONTENT_LEN], "doc_extract_failed"

    if ext == ".zip" or zipfile.is_zipfile(path):
        content, reason = extract_zip_summary(path)
        if content:
            return content, reason

    content = read_text_file(path)
    if content:
        return content, "text_fallback"
    return build_binary_file_summary(path, filename)


_USER_INTENT_LOOKBACK_TURNS = 3
_USER_INTENT_MAX_AGE_SECONDS = 20 * 60


def _recent_user_intent(user_id, now_ts: int | None = None) -> str:
    """Return the most recent non-empty user message prior to the current file.

    We look a few turns back (default 3) and only accept messages that are
    reasonably fresh (default 20 min). The goal is to capture intents like
    "帮我总结这个 pdf 的第三章" that arrive right before the attachment.
    """
    import time

    try:
        context = load_private_context(BASE_DATA_DIR, user_id)
    except Exception as e:
        print(f"[FILE] 读取私聊历史失败: {e}")
        return ""

    history = context.get("history") or []
    if not history:
        return ""

    cutoff = (now_ts or int(time.time())) - _USER_INTENT_MAX_AGE_SECONDS
    for item in reversed(history[-_USER_INTENT_LOOKBACK_TURNS:]):
        if not isinstance(item, dict):
            continue
        user_text = (item.get("user") or "").strip()
        if not user_text:
            continue
        ts = item.get("user_timestamp") or item.get("timestamp") or 0
        if ts and ts < cutoff:
            continue
        return user_text[:600]
    return ""


def handle_file_message(message_type, user_id, group_id, file_info):
    """Handle uploaded files for private or group contexts."""
    filename = derive_filename(file_info)
    file_url = file_info.get("url")
    file_path = file_info.get("path")
    safe_name = safe_filename(filename or "unknown_file")

    print(f"[FILE] 收到文件: name={filename!r}, url={file_url!r}, path={file_path!r}")

    if message_type == "private" and user_id != ALLOWED_PRIVATE_USER:
        print(f"[FILE] 非授权私聊用户 {user_id}，拒绝处理文件")
        return "ignore"

    if message_type == "group":
        send_group_msg(group_id, "为保护隐私，群聊模式下不会直接解析或输出文件内容，请改为私聊发送。")
        return {"status": "file_blocked_in_group"}

    if not filename:
        message_id = file_info.get("message_id")
        recovered_info = None
        if message_id:
            detail = get_msg_detail(message_id)
            if isinstance(detail, dict):
                recovered_info = extract_file_info(detail)
        if recovered_info:
            merged = file_info.copy()
            merged.update({k: v for k, v in recovered_info.items() if v not in (None, "", [])})
            file_info = merged
            filename = derive_filename(file_info)
            safe_name = safe_filename(filename or "unknown_file")
            print(
                f"[FILE] 通过 get_msg 补全文件信息成功: "
                f"name={file_info.get('name')!r}, uuid={file_info.get('uuid')!r}, url={file_info.get('url')!r}"
            )

    if not filename:
        msg = "已检测到文件事件，但暂时没拿到文件名。"
        send_private_msg(user_id, msg)
        return {"status": "file_no_name"}

    save_dir = PRIVATE_UPLOAD_DIR if message_type == "private" else GROUP_UPLOAD_DIR
    saved_path, reason = download_file_if_possible(file_info, save_dir)

    user_intent = _recent_user_intent(user_id) if message_type == "private" else ""
    if user_intent:
        print(f"[FILE] 捕获到用户最近意图: {user_intent[:80]}...")

    if not saved_path:
        msg = (
            f"已识别文件：{safe_name}\n当前未能获取可用下载链接，也无法从本地路径读取。\n"
            f"原因：{reason}\n请稍后重试，或检查 NapCat 文件接口配置。"
        )
        send_private_msg(user_id, msg)
        return {"status": "file_recognized_but_not_downloaded"}

    content, extract_reason = extract_file_content_for_ai(saved_path, safe_name)
    if not content:
        fallback_msg = f"文件已保存：{safe_name}\n但当前无法提取内容。"
        if user_intent:
            fallback_msg += f"\n(收到你之前的请求：{user_intent[:120]}，但缺少可读内容，无法执行)"
        send_private_msg(user_id, fallback_msg)
        return {"status": "file_read_failed"}

    if user_intent:
        task_line = f"用户在发送此文件之前提出的请求是：{user_intent}\n请优先按这个请求来处理文件；如果请求与文件无关，再退回到默认的文件说明模式。"
    else:
        task_line = (
            "用户没有附带文字说明。请默认做文件说明：告诉用户这是什么、主要内容、有哪些值得注意的信息。"
        )

    extract_note = ""
    if extract_reason == "doc_extract_failed":
        extract_note = (
            "\n注意：本地解析该文档失败（可能缺少 PyMuPDF / python-docx / openpyxl 依赖或文档是扫描件）。"
            "请基于下方的说明文本回应，并明确告诉用户你没能读到真实内容。"
        )

    query = (
        "你是 Candace 的文件阅读助手。"
        f"{task_line}"
        "如果文件本身无法完整转成纯文本，也要明确说明你是基于结构/元数据进行判断。"
        f"{extract_note}\n\n"
        f"文件名：{safe_name}\n保存路径：{saved_path}\n文件下载方式：{reason}\n文件内容提取方式：{extract_reason}\n\n"
        f"------- 文件内容开始 -------\n{content}\n------- 文件内容结束 -------"
    )
    reply = call_ai(query, metadata={"user_id": user_id, "prompt_mode": "file_understanding"})
    send_private_msg(user_id, reply)
    return {
        "status": "file_processed_private",
        "user_intent_used": bool(user_intent),
        "extract_reason": extract_reason,
    }
